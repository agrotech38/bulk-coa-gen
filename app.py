import streamlit as st
from docx import Document
import mammoth
import random
import pandas as pd
import io
import os
import zipfile
from datetime import datetime
import calendar

# ---------------------------
# Component Ranges (FROM FIRST CODE)
# ---------------------------
RANGES = {
    "fat": (0.45, 0.55),
    "air": (2.90, 3.10),
    "ash": (0.45, 0.55),
    "protein": (2.45, 2.55),
    "gum": (80.10, 89.95)
}

MIDS = {k: (v[0] + v[1]) / 2 for k, v in RANGES.items()}

# ---------------------------
# Distribution Engine (WATER FILLING)
# ---------------------------
def distribute_within_bounds(target, names, mins, maxs, weights):
    vals = {n: target * (weights[n] / sum(weights.values())) for n in names}
    locked = {n: False for n in names}

    for _ in range(100):
        for n in names:
            if not locked[n]:
                if vals[n] < mins[n]:
                    vals[n] = mins[n]
                    locked[n] = True
                if vals[n] > maxs[n]:
                    vals[n] = maxs[n]
                    locked[n] = True

        unlocked = [n for n in names if not locked[n]]
        if not unlocked:
            break

        remaining = target - sum(vals.values())
        wsum = sum(weights[n] for n in unlocked)

        for n in unlocked:
            vals[n] += remaining * (weights[n] / wsum)

    for n in names:
        vals[n] = round(vals[n], 2)

    diff = round(target - sum(vals.values()), 2)
    for n in names:
        if abs(diff) < 0.01:
            break
        low = RANGES[n][0]
        high = RANGES[n][1]
        if low <= vals[n] + diff <= high:
            vals[n] += diff
            break

    return vals

# ---------------------------
# RANDOM COMPONENT ENGINE
# ---------------------------
def calculate_components_random(moisture):
    remaining = round(100 - moisture, 4)
    others = ["fat", "air", "ash", "protein"]
    gum_min, gum_max = RANGES["gum"]

    for _ in range(2000):
        gum = round(random.uniform(gum_min, gum_max), 4)
        left = remaining - gum

        mins = {o: RANGES[o][0] for o in others}
        maxs = {o: RANGES[o][1] for o in others}
        weights = {o: random.random() + MIDS[o] for o in others}

        try:
            vals = distribute_within_bounds(left, others, mins, maxs, weights)
            total = moisture + gum + sum(vals.values())
            if abs(total - 100) <= 0.01:
                return (
                    round(gum, 2),
                    vals["protein"],
                    vals["ash"],
                    vals["air"],
                    vals["fat"],
                )
        except:
            pass

    raise ValueError("Cannot generate valid components")

# ---------------------------
# DOCX FUNCTIONS
# ---------------------------
def advanced_replace_text_preserving_style(doc, replacements):
    def replace_in_paragraph(paragraph):
        full = ''.join(r.text for r in paragraph.runs)
        for k, v in replacements.items():
            key = f"{{{{{k}}}}}"
            if key in full:
                paragraph.text = full.replace(key, v)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    replace_in_paragraph(p)

def generate_docx(data, template, out):
    doc = Document(template)
    advanced_replace_text_preserving_style(doc, data)
    doc.save(out)

# ---------------------------
# STREAMLIT UI
# ---------------------------
st.set_page_config("Bulk COA Generator", layout="wide")
st.title("📦 Bulk COA Generator (Scientific Randomization)")

uploaded_file = st.file_uploader("Upload Excel", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.success(f"Loaded {len(df)} rows")

    out_dir = "generated"
    os.makedirs(out_dir, exist_ok=True)
    files = []

    for i, row in df.iterrows():
        try:
            code = str(row["Code"])
            date = str(row["Date"])
            batch = str(row["Batch No"])
            moisture = float(row["Moisture"])
            ph = str(row["pH"])
            mesh = str(row["200 Mesh"])
            v2 = str(row["Viscosity 2H"])
            v24 = str(row["Viscosity 24H"])

            try:
                dt = datetime.strptime(date, "%B %Y")
                year = dt.year + 2
                month = dt.month - 1
                if month == 0:
                    month = 12
                    year -= 1
                best_before = f"{calendar.month_name[month].upper()} {year}"
            except:
                best_before = "N/A"

            gum, protein, ash, air, fat = calculate_components_random(moisture)

            data = {
                "DATE": date,
                "BATCH_NO": batch,
                "BEST_BEFORE": best_before,
                "MOISTURE": f"{moisture:.2f}%",
                "PH": ph,
                "MESH_200": f"{mesh}%",
                "VISCOSITY_2H": v2,
                "VISCOSITY_24H": v24,
                "GUM_CONTENT": f"{gum:.2f}%",
                "PROTEIN": f"{protein:.2f}%",
                "ASH_CONTENT": f"{ash:.2f}%",
                "AIR": f"{air:.2f}%",
                "FAT": f"{fat:.2f}%"
            }

            template = f"COA {code}.docx"
            safe = batch.replace("/", "_")
            out = f"{out_dir}/COA-{safe}-{code}.docx"

            if os.path.exists(template):
                generate_docx(data, template, out)
                files.append(out)
                st.success(f"Generated {os.path.basename(out)}")
            else:
                st.error(f"Missing template: {template}")

        except Exception as e:
            st.error(f"Row {i+1} failed: {e}")

    if files:
        zip_path = f"{out_dir}/ALL_COA.zip"
        with zipfile.ZipFile(zip_path, "w") as z:
            for f in files:
                z.write(f, os.path.basename(f))

        with open(zip_path, "rb") as f:
            st.download_button("📦 Download All COAs", f, file_name="ALL_COA.zip")
